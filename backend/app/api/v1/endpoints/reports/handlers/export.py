"""
报告导出API端点

提供报告导出为PDF、Word、图片等格式的功能

@author XieHe Medical System
@created 2025-09-24
"""

import os
import tempfile
import typing
from datetime import datetime
from typing import Any, Optional

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, Query, status
from fastapi.responses import StreamingResponse
from sqlalchemy import and_
from sqlalchemy.orm import Session
from starlette.background import BackgroundTask

from app.core.access.auth import get_current_active_user
from app.core.database.session import get_db
from app.core.system.concurrency import (
    ConcurrencyLimitExceeded,
    report_export_gate,
    require_report_export_slot,
)
from app.core.system.logger import LogLevel, logger
from app.models.report import DiagnosticReport

from ..schemas.export import (
    BatchExportRequest,
    ExportResponse,
)

router = APIRouter()


@router.post("/single", response_model=ExportResponse)
async def export_single_report(
    report_id: int,
    format: str = Query(..., pattern="^(pdf|word|image|html)$"),
    template: Optional[str] = Query(None),
    include_images: bool = Query(True),
    watermark: Optional[str] = Query(None),
    current_user: dict[str, typing.Any] = Depends(get_current_active_user),
    db: Session = Depends(get_db),
    _slot: None = Depends(require_report_export_slot),
) -> ExportResponse:
    """
    导出单个报告
    """
    try:
        # 获取报告
        report = (
            db.query(DiagnosticReport)
            .filter(
                and_(
                    DiagnosticReport.id == report_id,
                    DiagnosticReport.is_deleted.is_(False),
                )
            )
            .first()
        )

        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND, detail="报告不存在"
            )

        # 生成导出任务ID
        task_id = f"export_{report_id}_{format}_{int(datetime.now().timestamp())}"

        # 根据格式导出
        if format == "pdf":
            file_path = await export_to_pdf(report, template, include_images, watermark)
        elif format == "word":
            file_path = await export_to_word(
                report, template, include_images, watermark
            )
        elif format == "image":
            file_path = await export_to_image(
                report, template, include_images, watermark
            )
        elif format == "html":
            file_path = await export_to_html(
                report, template, include_images, watermark
            )
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST, detail="不支持的导出格式"
            )

        # 获取文件大小
        file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0

        logger.emit_event(
            LogLevel.INFO, message=f"报告导出成功: {report.report_number} -> {format}"
        )

        return ExportResponse(
            task_id=task_id,
            status="completed",
            message="导出成功",
            download_url=f"/api/v1/report-export/download/{task_id}",
            file_size=file_size,
            created_at=datetime.now(),
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.emit_event(LogLevel.ERROR, message=f"报告导出失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="报告导出失败"
        )


@router.post("/batch", response_model=ExportResponse)
async def export_batch_reports(
    request: BatchExportRequest,
    background_tasks: BackgroundTasks,
    current_user: dict[str, typing.Any] = Depends(get_current_active_user),
    db: Session = Depends(get_db),
    _slot: None = Depends(require_report_export_slot),
) -> ExportResponse:
    """
    批量导出报告
    """
    try:
        # 构建查询条件
        query = db.query(DiagnosticReport).filter(
            DiagnosticReport.is_deleted.is_(False)
        )

        # 应用筛选条件
        if request.filters.get("patient_id"):
            query = query.filter(
                DiagnosticReport.patient_id == request.filters["patient_id"]
            )
        if request.filters.get("status"):
            query = query.filter(DiagnosticReport.status == request.filters["status"])
        if request.filters.get("date_from"):
            query = query.filter(
                DiagnosticReport.report_date >= request.filters["date_from"]
            )
        if request.filters.get("date_to"):
            query = query.filter(
                DiagnosticReport.report_date <= request.filters["date_to"]
            )

        # 限制数量
        reports = query.limit(request.max_reports).all()

        if not reports:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND, detail="没有找到符合条件的报告"
            )

        # 生成批量导出任务ID
        task_id = f"batch_export_{request.format}_{int(datetime.now().timestamp())}"

        # 后台任务处理批量导出
        background_tasks.add_task(
            process_batch_export,
            task_id,
            reports,
            request.format,
            request.template,
            request.include_images,
            request.watermark,
        )

        logger.emit_event(
            LogLevel.INFO,
            message=f"批量导出任务创建: {task_id}, 报告数量: {len(reports)}",
        )

        return ExportResponse(
            task_id=task_id,
            status="processing",
            message=f"正在处理 {len(reports)} 个报告的批量导出",
            created_at=datetime.now(),
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.emit_event(LogLevel.ERROR, message=f"批量导出失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="批量导出失败"
        )


@router.get("/download/{task_id}")
async def download_exported_file(
    task_id: str, current_user: dict[str, typing.Any] = Depends(get_current_active_user)
) -> StreamingResponse:
    """
    下载导出的文件
    """
    try:
        # 查找导出文件
        file_path = f"/tmp/exports/{task_id}"

        # 检查文件是否存在
        if not os.path.exists(file_path):
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND, detail="导出文件不存在或已过期"
            )

        # 确定文件类型和名称
        if task_id.endswith(".pdf"):
            media_type = "application/pdf"
            filename = f"report_{task_id}.pdf"
        elif task_id.endswith(".docx"):
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"report_{task_id}.docx"
        elif task_id.endswith(".png"):
            media_type = "image/png"
            filename = f"report_{task_id}.png"
        elif task_id.endswith(".html"):
            media_type = "text/html"
            filename = f"report_{task_id}.html"
        else:
            media_type = "application/octet-stream"
            filename = f"report_{task_id}"

        logger.emit_event(LogLevel.INFO, message=f"下载导出文件: {filename}")

        file_handle = open(file_path, "rb")
        return StreamingResponse(
            file_handle,
            media_type=media_type,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
            background=BackgroundTask(file_handle.close),
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.emit_event(LogLevel.ERROR, message=f"文件下载失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="文件下载失败"
        )


@router.get("/status/{task_id}")
async def get_export_status(
    task_id: str, current_user: dict[str, typing.Any] = Depends(get_current_active_user)
) -> dict[str, typing.Any]:
    """
    获取导出任务状态
    """
    try:
        # 检查任务状态（这里简化处理，实际应该从数据库或缓存中获取）
        file_path = f"/tmp/exports/{task_id}"

        if os.path.exists(file_path):
            file_size = os.path.getsize(file_path)
            return {
                "task_id": task_id,
                "status": "completed",
                "message": "导出完成",
                "download_url": f"/api/v1/report-export/download/{task_id}",
                "file_size": file_size,
                "created_at": datetime.now(),
            }
        else:
            return {
                "task_id": task_id,
                "status": "processing",
                "message": "正在处理中...",
                "created_at": datetime.now(),
            }

    except Exception as e:
        logger.emit_event(LogLevel.ERROR, message=f"获取导出状态失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="获取导出状态失败"
        )


# 导出处理函数
async def export_to_pdf(
    report: typing.Any,
    template: typing.Any = None,
    include_images: typing.Any = True,
    watermark: typing.Any = None,
) -> typing.Any:
    """导出为PDF格式"""
    try:
        # 这里使用模拟实现，实际应该使用reportlab或weasyprint
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.platypus import Flowable, Paragraph, SimpleDocTemplate, Spacer

        # 创建临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")

        # 创建PDF文档
        doc = SimpleDocTemplate(temp_file.name, pagesize=A4)
        styles = getSampleStyleSheet()
        story: list[Flowable] = []

        # 添加标题
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # 居中
        )
        story.append(Paragraph(report.report_title or "医学影像诊断报告", title_style))
        story.append(Spacer(1, 12))

        # 添加报告内容
        content_style = styles["Normal"]

        if report.clinical_history:
            story.append(Paragraph("<b>临床病史:</b>", styles["Heading2"]))
            story.append(Paragraph(report.clinical_history, content_style))
            story.append(Spacer(1, 12))

        if report.examination_technique:
            story.append(Paragraph("<b>检查技术:</b>", styles["Heading2"]))
            story.append(Paragraph(report.examination_technique, content_style))
            story.append(Spacer(1, 12))

        if report.findings:
            story.append(Paragraph("<b>检查所见:</b>", styles["Heading2"]))
            story.append(Paragraph(report.findings, content_style))
            story.append(Spacer(1, 12))

        if report.impression:
            story.append(Paragraph("<b>诊断意见:</b>", styles["Heading2"]))
            story.append(Paragraph(report.impression, content_style))
            story.append(Spacer(1, 12))

        if report.recommendations:
            story.append(Paragraph("<b>建议:</b>", styles["Heading2"]))
            story.append(Paragraph(report.recommendations, content_style))
            story.append(Spacer(1, 12))

        # 添加水印
        if watermark:
            story.append(Spacer(1, 50))
            watermark_style = ParagraphStyle(
                "Watermark",
                parent=styles["Normal"],
                fontSize=10,
                textColor=colors.gray,
                alignment=1,
            )
            story.append(Paragraph(watermark, watermark_style))

        # 构建PDF
        doc.build(story)

        return temp_file.name

    except Exception as e:
        logger.emit_event(LogLevel.ERROR, message=f"PDF导出失败: {e}")
        raise Exception(f"PDF导出失败: {e}")


async def export_to_word(
    report: typing.Any,
    template: typing.Any = None,
    include_images: typing.Any = True,
    watermark: typing.Any = None,
) -> typing.Any:
    """导出为Word格式"""
    try:
        # 这里使用模拟实现，实际应该使用python-docx
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # 创建文档
        doc = Document()

        # 添加标题
        title = doc.add_heading(report.report_title or "医学影像诊断报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加报告内容
        if report.clinical_history:
            doc.add_heading("临床病史", level=1)
            doc.add_paragraph(report.clinical_history)

        if report.examination_technique:
            doc.add_heading("检查技术", level=1)
            doc.add_paragraph(report.examination_technique)

        if report.findings:
            doc.add_heading("检查所见", level=1)
            doc.add_paragraph(report.findings)

        if report.impression:
            doc.add_heading("诊断意见", level=1)
            doc.add_paragraph(report.impression)

        if report.recommendations:
            doc.add_heading("建议", level=1)
            doc.add_paragraph(report.recommendations)

        # 添加水印
        if watermark:
            doc.add_paragraph()
            watermark_p = doc.add_paragraph(watermark)
            watermark_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 保存到临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)

        return temp_file.name

    except Exception as e:
        logger.emit_event(LogLevel.ERROR, message=f"Word导出失败: {e}")
        raise Exception(f"Word导出失败: {e}")


async def export_to_image(
    report: typing.Any,
    template: typing.Any = None,
    include_images: typing.Any = True,
    watermark: typing.Any = None,
) -> typing.Any:
    """导出为图片格式"""
    try:
        # 这里使用模拟实现，实际应该使用PIL或其他图像库
        from PIL import Image, ImageDraw, ImageFont

        # 创建图片
        width, height = 800, 1200
        img = Image.new("RGB", (width, height), color="white")
        draw = ImageDraw.Draw(img)

        # 设置字体（这里使用默认字体）
        font_title: ImageFont.ImageFont | ImageFont.FreeTypeFont
        font_heading: ImageFont.ImageFont | ImageFont.FreeTypeFont
        font_content: ImageFont.ImageFont | ImageFont.FreeTypeFont
        try:
            font_title = ImageFont.truetype("arial.ttf", 24)
            font_heading = ImageFont.truetype("arial.ttf", 18)
            font_content = ImageFont.truetype("arial.ttf", 14)
        except OSError:
            font_title = ImageFont.load_default()
            font_heading = ImageFont.load_default()
            font_content = ImageFont.load_default()

        y_position = 50

        # 添加标题
        title = report.report_title or "医学影像诊断报告"
        draw.text(
            (width // 2 - len(title) * 6, y_position),
            title,
            fill="black",
            font=font_title,
        )
        y_position += 80

        # 添加内容
        sections = [
            ("临床病史", report.clinical_history),
            ("检查技术", report.examination_technique),
            ("检查所见", report.findings),
            ("诊断意见", report.impression),
            ("建议", report.recommendations),
        ]

        for section_title, content in sections:
            if content:
                # 添加段落标题
                draw.text(
                    (50, y_position), section_title, fill="black", font=font_heading
                )
                y_position += 30

                # 添加内容（简单换行处理）
                lines = content.split("\n")
                for line in lines:
                    if y_position > height - 100:
                        break
                    draw.text(
                        (50, y_position), line[:80], fill="black", font=font_content
                    )
                    y_position += 20

                y_position += 20

        # 添加水印
        if watermark and y_position < height - 50:
            draw.text(
                (width // 2 - len(watermark) * 3, height - 50),
                watermark,
                fill="gray",
                font=font_content,
            )

        # 保存到临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        img.save(temp_file.name, "PNG")

        return temp_file.name

    except Exception as e:
        logger.emit_event(LogLevel.ERROR, message=f"图片导出失败: {e}")
        raise Exception(f"图片导出失败: {e}")


async def export_to_html(
    report: typing.Any,
    template: typing.Any = None,
    include_images: typing.Any = True,
    watermark: typing.Any = None,
) -> typing.Any:
    """导出为HTML格式"""
    try:
        # HTML模板
        html_template = """
        <!DOCTYPE html>
        <html lang="zh-CN">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{title}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                .header {{ text-align: center; margin-bottom: 40px; }}
                .title {{ font-size: 24px; font-weight: bold; margin-bottom: 20px; }}
                .section {{ margin-bottom: 30px; }}
                .section-title {{ font-size: 18px; font-weight: bold; margin-bottom: 10px; color: #333; }}
                .content {{ margin-left: 20px; }}
                .watermark {{ text-align: center; color: #999; font-size: 12px; margin-top: 50px; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="title">{title}</div>
            </div>

            {sections}

            {watermark_html}
        </body>
        </html>
        """

        # 构建段落HTML
        sections_html = ""
        sections = [
            ("临床病史", report.clinical_history),
            ("检查技术", report.examination_technique),
            ("检查所见", report.findings),
            ("诊断意见", report.impression),
            ("建议", report.recommendations),
        ]

        for section_title, content in sections:
            if content:
                sections_html += f"""
                <div class="section">
                    <div class="section-title">{section_title}</div>
                    <div class="content">{content.replace(chr(10), "<br>")}</div>
                </div>
                """

        # 水印HTML
        watermark_html = (
            f'<div class="watermark">{watermark}</div>' if watermark else ""
        )

        # 生成完整HTML
        html_content = html_template.format(
            title=report.report_title or "医学影像诊断报告",
            sections=sections_html,
            watermark_html=watermark_html,
        )

        # 保存到临时文件
        temp_file = tempfile.NamedTemporaryFile(
            delete=False, suffix=".html", mode="w", encoding="utf-8"
        )
        temp_file.write(html_content)
        temp_file.close()

        return temp_file.name

    except Exception as e:
        logger.emit_event(LogLevel.ERROR, message=f"HTML导出失败: {e}")
        raise Exception(f"HTML导出失败: {e}")


# 后台任务处理函数
async def process_batch_export(
    task_id: typing.Any,
    reports: typing.Any,
    format: typing.Any,
    template: typing.Any,
    include_images: typing.Any,
    watermark: typing.Any,
) -> None:
    """处理批量导出任务"""
    try:
        async with report_export_gate.acquire():
            await _process_batch_export(
                task_id, reports, format, template, include_images, watermark
            )
    except ConcurrencyLimitExceeded:
        logger.emit_event(
            LogLevel.WARNING, message=f"批量导出并发已满，跳过任务: {task_id}"
        )


async def _export_report(
    report: Any,
    export_format: str,
    template: Any,
    include_images: bool,
    watermark: str | None,
) -> str:
    exporters = {
        "pdf": export_to_pdf,
        "word": export_to_word,
        "image": export_to_image,
        "html": export_to_html,
    }
    exporter = exporters.get(export_format)
    if exporter is None:
        raise ValueError(f"不支持的导出格式: {export_format}")
    return str(await exporter(report, template, include_images, watermark))


def _package_exported_files(
    task_id: str,
    export_format: str,
    export_dir: str,
    exported_files: list[str],
) -> str:
    if not exported_files:
        raise RuntimeError("没有成功导出任何报告")
    if len(exported_files) == 1:
        final_path = f"{export_dir}/{task_id}.{export_format}"
        os.rename(exported_files[0], final_path)
        return final_path

    import zipfile

    final_path = f"{export_dir}/{task_id}.zip"
    with zipfile.ZipFile(final_path, "w") as archive:
        for index, file_path in enumerate(exported_files, start=1):
            archive.write(file_path, f"report_{index}.{export_format}")
            os.unlink(file_path)
    return final_path


async def _process_batch_export(
    task_id: typing.Any,
    reports: typing.Any,
    format: typing.Any,
    template: typing.Any,
    include_images: typing.Any,
    watermark: typing.Any,
) -> None:
    """Run the actual batch export work after the caller has acquired a slot."""
    try:
        logger.emit_event(LogLevel.INFO, message=f"开始批量导出任务: {task_id}")

        # 创建导出目录
        export_dir = "/tmp/exports"
        os.makedirs(export_dir, exist_ok=True)

        exported_files: list[str] = []

        # 逐个导出报告
        for report in reports:
            try:
                exported_files.append(
                    await _export_report(
                        report,
                        str(format),
                        template,
                        bool(include_images),
                        str(watermark) if watermark is not None else None,
                    )
                )

            except Exception as e:
                logger.emit_event(
                    LogLevel.ERROR, message=f"报告 {report.id} 导出失败: {e}"
                )
                continue

        final_path = _package_exported_files(
            str(task_id), str(format), export_dir, exported_files
        )

        logger.emit_event(
            LogLevel.INFO, message=f"批量导出任务完成: {task_id}, 文件: {final_path}"
        )

    except Exception as e:
        logger.emit_event(
            LogLevel.ERROR, message=f"批量导出任务失败: {task_id}, 错误: {e}"
        )
        raise
